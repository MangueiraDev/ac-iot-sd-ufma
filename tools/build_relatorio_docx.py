import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / ".codex_docx_deps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import html as lxml_html
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "RELATORIO_TECNICO_FINAL_DOCX_SOURCE.html"
OUT = ROOT / "docs" / "RELATORIO_TECNICO_FINAL.docx"

BLUE = RGBColor(16, 32, 57)
MUTED = RGBColor(81, 98, 122)
BORDER = "D7DFEB"
HEAD_FILL = "EAF1FB"
ABSTRACT_FILL = "F1F6FF"
CODE_FILL = "0B1220"
CODE_TEXT = RGBColor(234, 242, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, pct=5000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct))


def set_paragraph_border_bottom(paragraph, color=BORDER):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def text_content(node):
    return " ".join(" ".join(node.itertext()).split())


def add_runs_from_node(paragraph, node):
    if node.text:
        paragraph.add_run(node.text)
    for child in node:
        run_text = child.text_content() if child.tag in ("strong", "em", "code") else None
        if run_text is not None:
            run = paragraph.add_run(run_text)
            if child.tag == "strong":
                run.bold = True
            if child.tag == "em":
                run.italic = True
            if child.tag == "code":
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            if child.tail:
                paragraph.add_run(child.tail)
        else:
            add_runs_from_node(paragraph, child)
            if child.tail:
                paragraph.add_run(child.tail)


def add_paragraph(doc, node=None, text=None, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    if node is not None:
        add_runs_from_node(p, node)
    elif text:
        p.add_run(text)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = BLUE
        set_paragraph_border_bottom(p)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    else:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12.5)
        run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(5)
    return p


def add_html_table(doc, table_node):
    rows = table_node.xpath("./tr")
    if not rows:
        return
    col_count = max(len(r.xpath("./th|./td")) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    is_toc = "toc" in (table_node.get("class") or "")
    for r_idx, row_node in enumerate(rows):
        cells = row_node.xpath("./th|./td")
        for c_idx, cell_node in enumerate(cells):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, color="FFFFFF" if is_toc else BORDER, size="2" if is_toc else "6")
            if cell_node.tag == "th":
                set_cell_shading(cell, HEAD_FILL)
            text = text_content(cell_node)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if is_toc and c_idx == col_count - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text)
            run.font.size = Pt(9.3 if not is_toc else 10.5)
            if cell_node.tag == "th":
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def image_width_for(path):
    with Image.open(path) as img:
        w, h = img.size
    max_width = Inches(6.35)
    if w >= h:
        return max_width
    return Inches(4.8)


def add_figure(doc, figure_node):
    img = figure_node.xpath(".//img")
    cap = figure_node.xpath('.//*[contains(@class, "caption")]')
    if img:
        src = img[0].get("src")
        path = ROOT / "docs" / src
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if path.exists():
            run.add_picture(str(path), width=image_width_for(path))
    if cap:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text_content(cap[0]))
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED


def add_code(doc, node):
    text = node.text_content()
    p = doc.add_paragraph()
    set_paragraph_shading(p, CODE_FILL)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = CODE_TEXT


def build_doc():
    tree = lxml_html.fromstring(SRC.read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(7)

    # Cover
    cover = tree.xpath('//section[contains(@class, "cover")]')[0]
    for cls, size, bold, before, after in [
        ("institution", 11, True, 0, 4),
        (None, 11, False, 0, 4),
        ("course", 11, True, 0, 120),
    ]:
        if cls:
            text = text_content(cover.xpath(f'.//*[contains(@class, "{cls}")]')[0])
        else:
            text = text_content(cover.xpath("./div[1]/div[2]")[0])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    title_run = p.add_run("Relatório Técnico Final")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = BLUE
    title_run.add_break()
    title_run = p.add_run("Gestão Energética dos Ar-Condicionados das Salas de Aula do Prédio Paulo Freire via IoT")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(45)
    subtitle_run = p.add_run(text_content(cover.xpath('.//*[contains(@class, "subtitle")]')[0]))
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(49, 68, 95)

    team_text = cover.xpath('.//*[contains(@class, "team")]')[0].text_content()
    for line in [ln.strip() for ln in team_text.splitlines() if ln.strip()]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1.45)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        if line.endswith(":"):
            run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.add_run("São Luís - MA").add_break()
    p.add_run("2026")
    doc.add_page_break()

    # Body, skip cover section but keep all others flattened.
    for section_node in tree.xpath("//section[not(contains(@class, 'cover'))]"):
        for node in section_node:
            if not isinstance(node.tag, str):
                continue
            classes = node.get("class") or ""
            if node.tag == "h2":
                if "toc-heading" in classes or node.get("id") == "referencias":
                    if len(doc.paragraphs) > 0:
                        doc.add_page_break()
                add_heading(doc, text_content(node), 2)
            elif node.tag == "h3":
                add_heading(doc, text_content(node), 3)
            elif node.tag == "p":
                add_paragraph(doc, node=node)
            elif node.tag == "div" and "abstract" in classes:
                for p_node in node.xpath("./p"):
                    p = add_paragraph(doc, node=p_node)
                    set_paragraph_shading(p, ABSTRACT_FILL)
                for div_node in node.xpath("./div"):
                    p = add_paragraph(doc, node=div_node)
                    for run in p.runs:
                        run.bold = True
            elif node.tag == "table":
                add_html_table(doc, node)
            elif node.tag == "ol":
                for li in node.xpath("./li"):
                    p = doc.add_paragraph(style="List Number")
                    add_runs_from_node(p, li)
            elif node.tag == "ul":
                for li in node.xpath("./li"):
                    p = doc.add_paragraph(style="List Bullet")
                    add_runs_from_node(p, li)
            elif node.tag == "div" and "figure" in classes:
                add_figure(doc, node)
            elif node.tag == "pre":
                add_code(doc, node)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
